from docx import Document

class DOCXReader:
    """
    DOCX parser that extracts:
    - Paragraphs
    - Tables (as structured rows)
    """

    def load(self, file_path: str):
        doc = Document(file_path)
        results = []

        # Paragraphs
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text:
                results.append({
                    "text": text,
                    "source": "docx",
                    "type": "paragraph",
                    "metadata": {
                        "paragraph_index": i
                    }
                })

        # Tables
        for tbl_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                row_text = " | ".join([c.text.strip() for c in row.cells])
                results.append({
                    "text": row_text,
                    "source": "docx",
                    "type": "table_row",
                    "metadata": {
                        "table_index": tbl_idx,
                        "row_index": row_idx
                    }
                })

        return results
